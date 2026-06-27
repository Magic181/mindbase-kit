import base64
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.test import TestCase, TransactionTestCase, override_settings
from docx import Document as DocxDocument
from rest_framework import status
from rest_framework.test import APIClient

from apps.notebooks.models import Notebook

from .assets import extract_assets
from .chunking import chunk_blocks
from .models import Document, DocumentAsset, DocumentChunk, DocumentStatus
from .parsers import _docx_heading_level, parse_file, parse_file_blocks
from .serializers import DocumentSerializer
from .tasks import parse_document_task
from .vision import _call_vision_model


TINY_PNG = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII='
)


class FakeResponse:
    def __init__(self, status_code=200, payload=None, text=''):
        self.status_code = status_code
        self._payload = payload or {}
        self.text = text

    def json(self):
        return self._payload


class DocumentUploadTests(TransactionTestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.settings_override = override_settings(
            MEDIA_ROOT=Path(self.temp_dir.name),
            MAX_UPLOAD_SIZE_MB=20,
        )
        self.settings_override.enable()

        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username='upload-user',
            email='upload@example.com',
            password='test-password',
        )
        self.notebook = Notebook.objects.create(
            user=self.user,
            name='Upload notebook',
        )
        self.client = APIClient()
        self.client.force_authenticate(self.user)

    def tearDown(self):
        self.settings_override.disable()
        self.temp_dir.cleanup()

    def test_rejects_mixed_invalid_upload_before_saving_any_file(self):
        valid_file = self._file('valid.txt', b'valid content')
        invalid_file = self._file('slides.pptx', b'not supported')

        with patch('apps.documents.views.save_uploaded_file') as save_uploaded_file:
            response = self.client.post(
                f'/api/v1/notebooks/{self.notebook.id}/documents/',
                {'files': [valid_file, invalid_file]},
                format='multipart',
            )

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertFalse(save_uploaded_file.called)
        self.assertEqual(Document.objects.count(), 0)

    def test_valid_upload_creates_document_and_enqueues_parse_after_commit(self):
        uploaded_file = self._file('notes.md', b'# Notes\n\nBatch upload test.')

        with patch('apps.documents.views.enqueue_parse_task') as enqueue_parse_task:
            response = self.client.post(
                f'/api/v1/notebooks/{self.notebook.id}/documents/',
                {'files': [uploaded_file]},
                format='multipart',
            )

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        document = Document.objects.get()
        self.assertEqual(document.name, 'notes.md')
        self.assertEqual(document.file_type, 'md')
        self.assertEqual(document.status, DocumentStatus.UPLOADING)
        self.assertTrue((Path(self.temp_dir.name) / document.file_path).exists())
        enqueue_parse_task.assert_called_once_with(document.id)

    @patch('apps.documents.views.enqueue_parse_task')
    def test_reparse_completed_document_enqueues_parse(self, enqueue_parse_task):
        document = Document.objects.create(
            notebook=self.notebook,
            name='notes.md',
            file_path='files/notes.md',
            file_type='md',
            status=DocumentStatus.COMPLETED,
            chunk_count=3,
            error_message='old error',
        )

        response = self.client.post(f'/api/v1/documents/{document.id}/reparse/')

        self.assertEqual(response.status_code, status.HTTP_202_ACCEPTED)
        enqueue_parse_task.assert_called_once_with(document.id)
        document.refresh_from_db()
        self.assertEqual(document.status, DocumentStatus.PARSING)
        self.assertEqual(document.chunk_count, 0)
        self.assertEqual(document.error_message, '')

    @patch('apps.documents.views.enqueue_parse_task')
    def test_reparse_rejects_active_document(self, enqueue_parse_task):
        document = Document.objects.create(
            notebook=self.notebook,
            name='notes.md',
            file_path='files/notes.md',
            file_type='md',
            status=DocumentStatus.PARSING,
        )

        response = self.client.post(f'/api/v1/documents/{document.id}/reparse/')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        enqueue_parse_task.assert_not_called()

    @staticmethod
    def _file(name: str, content: bytes):
        from django.core.files.uploadedfile import SimpleUploadedFile

        return SimpleUploadedFile(name, content)


class DocxParsingTests(TestCase):
    def test_docx_parser_preserves_paragraph_table_order(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'report.docx'
            self._write_docx_with_table(path)

            blocks = parse_file_blocks(path, 'docx')
            text = parse_file(path, 'docx')

        self.assertEqual(
            [block['metadata']['source_type'] for block in blocks],
            ['paragraph', 'table', 'paragraph'],
        )
        self.assertIn('实验目标', blocks[0]['content'])
        self.assertEqual(blocks[0]['metadata']['file_type'], 'docx')
        self.assertEqual(blocks[0]['metadata']['parser_version'], 1)
        self.assertIn('模块 | 状态', blocks[1]['content'])
        self.assertEqual(blocks[1]['metadata']['file_type'], 'docx')
        self.assertIn('行 2: 登录 | 通过', blocks[1]['content'])
        self.assertIn('实验结论', blocks[2]['content'])
        self.assertLess(text.index('实验目标'), text.index('[表格 1]'))
        self.assertLess(text.index('[表格 1]'), text.index('实验结论'))

    def test_chunk_blocks_keeps_table_metadata(self):
        chunks = chunk_blocks([
            {
                'content': '实验目标',
                'metadata': {'source_type': 'paragraph', 'block_index': 0},
            },
            {
                'content': '[表格 1]\n行 1: 模块 | 状态',
                'metadata': {
                    'source_type': 'table',
                    'block_index': 1,
                    'table_index': 1,
                    'row_count': 1,
                    'col_count': 2,
                },
            },
        ])

        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0]['metadata']['source_type'], 'paragraph')
        self.assertEqual(chunks[1]['metadata']['source_type'], 'table')
        self.assertEqual(chunks[1]['metadata']['table_index'], 1)

    def test_chunk_blocks_preserves_file_type_when_merging_blocks(self):
        chunks = chunk_blocks([
            {
                'content': '第一段',
                'metadata': {
                    'source_type': 'paragraph',
                    'block_index': 0,
                    'file_type': 'md',
                    'parser_version': 1,
                },
            },
            {
                'content': '第二段',
                'metadata': {
                    'source_type': 'paragraph',
                    'block_index': 1,
                    'file_type': 'md',
                    'parser_version': 1,
                },
            },
        ])

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0]['metadata']['source_type'], 'paragraph')
        self.assertEqual(chunks[0]['metadata']['file_type'], 'md')
        self.assertEqual(chunks[0]['metadata']['block_indexes'], [0, 1])

    def test_text_parser_uses_unified_block_metadata(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'notes.txt'
            path.write_text('第一段\n\n第二段', encoding='utf-8')

            blocks = parse_file_blocks(path, 'txt')

        self.assertEqual(len(blocks), 2)
        self.assertEqual(blocks[0]['source_type'], 'paragraph')
        self.assertEqual(blocks[0]['metadata']['file_type'], 'txt')
        self.assertEqual(blocks[0]['metadata']['parser_version'], 1)

    def test_markdown_parser_extracts_heading_code_and_table_blocks(self):
        markdown = (
            '# 实验报告\n\n'
            '正文段落。\n\n'
            '| 模块 | 状态 |\n'
            '| --- | --- |\n'
            '| 登录 | 通过 |\n\n'
            '```python\n'
            'print("ok")\n'
            '```\n'
        )
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'report.md'
            path.write_text(markdown, encoding='utf-8')

            blocks = parse_file_blocks(path, 'md')

        self.assertEqual(
            [block['metadata']['source_type'] for block in blocks],
            ['heading', 'paragraph', 'table', 'code'],
        )
        self.assertEqual(blocks[0]['metadata']['heading_level'], 1)
        self.assertEqual(blocks[0]['metadata']['file_type'], 'md')
        self.assertIn('[表格 1]', blocks[2]['content'])
        self.assertIn('行 1: 模块 | 状态', blocks[2]['content'])
        self.assertIn('行 2: 登录 | 通过', blocks[2]['content'])
        self.assertEqual(blocks[2]['metadata']['table_index'], 1)
        self.assertEqual(blocks[2]['metadata']['row_count'], 2)
        self.assertEqual(blocks[2]['metadata']['col_count'], 2)
        self.assertEqual(blocks[3]['metadata']['language'], 'python')

    def test_docx_parser_detects_heading_style(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'heading.docx'
            doc = DocxDocument()
            paragraph = doc.add_paragraph('章节标题')
            paragraph.style = 'Heading 1'
            doc.save(path)

            blocks = parse_file_blocks(path, 'docx')

        self.assertEqual(blocks[0]['metadata']['source_type'], 'heading')
        self.assertEqual(blocks[0]['metadata']['heading_level'], 1)

    def test_docx_heading_level_supports_style_variants(self):
        self.assertEqual(_docx_heading_level('Heading1'), 1)
        self.assertEqual(_docx_heading_level('Heading 2'), 2)
        self.assertEqual(_docx_heading_level('标题1'), 1)
        self.assertEqual(_docx_heading_level('标题 2'), 2)

    def test_markdown_asset_extraction_registers_image_reference(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'report.md'
            image_path = Path(temp_dir) / 'diagram.png'
            image_path.write_bytes(TINY_PNG)
            path.write_text('前文说明\n\n![流程图](diagram.png)\n\n后文说明', encoding='utf-8')

            assets = extract_assets(path, 'md', [])

        self.assertEqual(len(assets), 1)
        self.assertEqual(assets[0].original_name, 'diagram.png')
        self.assertIn('前文说明', assets[0].nearby_text)
        self.assertEqual(assets[0].metadata['alt_text'], '流程图')
        self.assertTrue(assets[0].metadata['is_local'])

    def test_docx_asset_extraction_reads_embedded_images(self):
        with TemporaryDirectory() as temp_dir:
            image_path = Path(temp_dir) / 'pixel.png'
            image_path.write_bytes(TINY_PNG)
            docx_path = Path(temp_dir) / 'image.docx'
            doc = DocxDocument()
            doc.add_paragraph('图片说明：系统流程图。')
            doc.add_picture(str(image_path))
            doc.save(docx_path)

            blocks = parse_file_blocks(docx_path, 'docx')
            assets = extract_assets(docx_path, 'docx', blocks)

        self.assertEqual(len(assets), 1)
        self.assertEqual(assets[0].metadata['source'], 'docx')
        self.assertIn('图片说明', assets[0].nearby_text)
        self.assertTrue(assets[0].data)

    @patch('apps.documents.parsers.PdfReader')
    def test_pdf_parser_uses_page_blocks_with_unified_metadata(self, pdf_reader):
        class FakePage:
            def __init__(self, text):
                self.text = text

            def extract_text(self):
                return self.text

        pdf_reader.return_value.pages = [FakePage('第一页内容'), FakePage('第二页内容')]

        blocks = parse_file_blocks(Path('fake.pdf'), 'pdf')

        self.assertEqual([block['metadata']['source_type'] for block in blocks], ['page', 'page'])
        self.assertEqual(blocks[0]['metadata']['file_type'], 'pdf')
        self.assertEqual(blocks[0]['metadata']['page'], 1)
        self.assertIn('[第2页]', blocks[1]['content'])

    @staticmethod
    def _write_docx_with_table(path: Path) -> None:
        doc = DocxDocument()
        doc.add_paragraph('实验目标：验证 DOCX 表格解析。')
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '模块'
        table.cell(0, 1).text = '状态'
        table.cell(1, 0).text = '登录'
        table.cell(1, 1).text = '通过'
        doc.add_paragraph('实验结论：表格内容已被读取。')
        doc.save(path)


class VisionProviderTests(TestCase):
    @patch.dict('os.environ', {
        'VISION_PROVIDER': 'openai_compatible',
        'VISION_BASE_URL': 'https://vision.example/v1',
        'VISION_MODEL': 'test-vlm',
        'VISION_TIMEOUT_SECONDS': '10',
    }, clear=True)
    @patch('apps.documents.vision.requests.post')
    def test_openai_compatible_provider_sends_data_url(self, post):
        post.return_value = FakeResponse(payload={
            'choices': [
                {'message': {'content': '图片描述'}}
            ]
        })
        with TemporaryDirectory() as temp_dir:
            image_path = Path(temp_dir) / 'image.png'
            image_path.write_bytes(TINY_PNG)
            asset = DocumentAsset(
                original_name='image.png',
                nearby_text='附近说明',
            )

            result = _call_vision_model(asset, image_path, 'test-key')

        self.assertEqual(result['status'], DocumentAsset.VisionStatus.SUCCESS)
        self.assertEqual(result['provider'], 'openai_compatible')
        post.assert_called_once()
        url = post.call_args.args[0]
        payload = post.call_args.kwargs['json']
        self.assertEqual(url, 'https://vision.example/v1/chat/completions')
        image_url = payload['messages'][0]['content'][1]['image_url']['url']
        self.assertTrue(image_url.startswith('data:image/png;base64,'))
        self.assertNotIn('thinking', payload)

    @patch.dict('os.environ', {
        'VISION_PROVIDER': 'zhipu',
        'VISION_MODEL': 'glm-4.6v-flash',
        'VISION_TIMEOUT_SECONDS': '10',
        'VISION_THINKING': 'disabled',
    }, clear=True)
    @patch('apps.documents.vision.requests.post')
    def test_zhipu_provider_sends_raw_base64_and_thinking(self, post):
        post.return_value = FakeResponse(payload={
            'choices': [
                {'message': {'content': '智谱图片描述'}}
            ]
        })
        with TemporaryDirectory() as temp_dir:
            image_path = Path(temp_dir) / 'image.png'
            image_path.write_bytes(TINY_PNG)
            asset = DocumentAsset(original_name='image.png')

            result = _call_vision_model(asset, image_path, 'test-key')

        self.assertEqual(result['status'], DocumentAsset.VisionStatus.SUCCESS)
        self.assertEqual(result['provider'], 'zhipu')
        post.assert_called_once()
        url = post.call_args.args[0]
        payload = post.call_args.kwargs['json']
        image_url = payload['messages'][0]['content'][1]['image_url']['url']
        self.assertEqual(url, 'https://open.bigmodel.cn/api/paas/v4/chat/completions')
        self.assertFalse(image_url.startswith('data:'))
        self.assertEqual(base64.b64decode(image_url), TINY_PNG)
        self.assertEqual(payload['thinking'], {'type': 'disabled'})


class ParseDocumentTaskTests(TransactionTestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.settings_override = override_settings(MEDIA_ROOT=Path(self.temp_dir.name))
        self.settings_override.enable()

        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username='parse-user',
            email='parse@example.com',
            password='test-password',
        )
        self.notebook = Notebook.objects.create(user=self.user, name='Parse notebook')

    def tearDown(self):
        self.settings_override.disable()
        self.temp_dir.cleanup()

    def test_parse_task_stores_docx_table_chunks_with_metadata(self):
        relative_path = Path('files') / str(self.user.id) / str(self.notebook.id) / 'report.docx'
        file_path = Path(self.temp_dir.name) / relative_path
        file_path.parent.mkdir(parents=True, exist_ok=True)
        DocxParsingTests._write_docx_with_table(file_path)

        document = Document.objects.create(
            notebook=self.notebook,
            name='report.docx',
            file_path=str(relative_path),
            file_size=file_path.stat().st_size,
            file_type='docx',
            status=DocumentStatus.UPLOADING,
        )

        parse_document_task(document.id)

        document.refresh_from_db()
        self.assertEqual(document.status, DocumentStatus.COMPLETED)
        self.assertGreaterEqual(document.chunk_count, 2)
        table_chunk = DocumentChunk.objects.get(
            document=document,
            metadata__source_type='table',
        )
        self.assertIn('模块 | 状态', table_chunk.content)
        self.assertEqual(table_chunk.metadata['table_index'], 1)
        self.assertEqual(table_chunk.metadata['file_type'], 'docx')

    def test_parse_task_stores_markdown_image_assets(self):
        relative_dir = Path('files') / str(self.user.id) / str(self.notebook.id)
        file_path = Path(self.temp_dir.name) / relative_dir / 'report.md'
        image_path = Path(self.temp_dir.name) / relative_dir / 'diagram.png'
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text('# 报告\n\n![流程图](diagram.png)', encoding='utf-8')
        image_path.write_bytes(TINY_PNG)

        document = Document.objects.create(
            notebook=self.notebook,
            name='report.md',
            file_path=str(relative_dir / 'report.md'),
            file_size=file_path.stat().st_size,
            file_type='md',
            status=DocumentStatus.UPLOADING,
        )

        parse_document_task(document.id)

        document.refresh_from_db()
        self.assertEqual(document.status, DocumentStatus.COMPLETED)
        asset = DocumentAsset.objects.get(document=document)
        self.assertEqual(asset.asset_type, DocumentAsset.AssetType.IMAGE)
        self.assertEqual(asset.original_name, 'diagram.png')
        self.assertTrue(asset.file_path)
        self.assertTrue((Path(self.temp_dir.name) / asset.file_path).exists())
        self.assertEqual(asset.metadata['alt_text'], '流程图')

    @patch('apps.documents.ocr.run_asset_ocr')
    def test_parse_task_creates_image_ocr_chunks_when_ocr_succeeds(self, run_asset_ocr):
        run_asset_ocr.return_value = {
            'status': DocumentAsset.OCRStatus.SUCCESS,
            'text': '流程图包含上传、解析、检索三个步骤。',
            'error': '',
            'engine': 'test-ocr',
        }
        relative_dir = Path('files') / str(self.user.id) / str(self.notebook.id)
        file_path = Path(self.temp_dir.name) / relative_dir / 'report.md'
        image_path = Path(self.temp_dir.name) / relative_dir / 'diagram.png'
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text('# 报告\n\n![流程图](diagram.png)', encoding='utf-8')
        image_path.write_bytes(TINY_PNG)

        document = Document.objects.create(
            notebook=self.notebook,
            name='report.md',
            file_path=str(relative_dir / 'report.md'),
            file_size=file_path.stat().st_size,
            file_type='md',
            status=DocumentStatus.UPLOADING,
        )

        parse_document_task(document.id)

        document.refresh_from_db()
        asset = DocumentAsset.objects.get(document=document)
        self.assertEqual(asset.ocr_status, DocumentAsset.OCRStatus.SUCCESS)
        self.assertIn('上传、解析、检索', asset.ocr_text)

        ocr_chunk = DocumentChunk.objects.get(
            document=document,
            metadata__source_type='image_ocr',
        )
        self.assertIn('OCR 文本', ocr_chunk.content)
        self.assertIn('上传、解析、检索', ocr_chunk.content)
        self.assertEqual(ocr_chunk.metadata['asset_id'], asset.id)
        self.assertEqual(ocr_chunk.metadata['ocr_engine'], 'test-ocr')
        self.assertEqual(document.chunk_count, document.chunks.count())

    @patch('apps.documents.vision.run_asset_vision')
    @patch('apps.documents.ocr.run_asset_ocr')
    def test_parse_task_creates_image_caption_chunks_when_vision_succeeds(
        self,
        run_asset_ocr,
        run_asset_vision,
    ):
        run_asset_ocr.return_value = {
            'status': DocumentAsset.OCRStatus.SKIPPED,
            'text': '',
            'error': 'No OCR text extracted',
        }
        run_asset_vision.return_value = {
            'status': DocumentAsset.VisionStatus.SUCCESS,
            'text': '这是一张系统流程图，展示上传、解析、检索三个阶段。',
            'error': '',
            'model': 'test-vision',
        }
        relative_dir = Path('files') / str(self.user.id) / str(self.notebook.id)
        file_path = Path(self.temp_dir.name) / relative_dir / 'report.md'
        image_path = Path(self.temp_dir.name) / relative_dir / 'diagram.png'
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text('# 报告\n\n![流程图](diagram.png)', encoding='utf-8')
        image_path.write_bytes(TINY_PNG)

        document = Document.objects.create(
            notebook=self.notebook,
            name='report.md',
            file_path=str(relative_dir / 'report.md'),
            file_size=file_path.stat().st_size,
            file_type='md',
            status=DocumentStatus.UPLOADING,
        )

        parse_document_task(document.id)

        document.refresh_from_db()
        asset = DocumentAsset.objects.get(document=document)
        self.assertEqual(asset.vision_status, DocumentAsset.VisionStatus.SUCCESS)
        self.assertIn('系统流程图', asset.vision_text)

        caption_chunk = DocumentChunk.objects.get(
            document=document,
            metadata__source_type='image_caption',
        )
        self.assertIn('视觉描述', caption_chunk.content)
        self.assertIn('上传、解析、检索', caption_chunk.content)
        self.assertEqual(caption_chunk.metadata['asset_id'], asset.id)
        self.assertEqual(caption_chunk.metadata['vision_model'], 'test-vision')
        self.assertEqual(document.chunk_count, document.chunks.count())

    def test_document_serializer_includes_asset_count(self):
        document = Document.objects.create(
            notebook=self.notebook,
            name='with-image.md',
            file_path='files/with-image.md',
            file_type='md',
            status=DocumentStatus.COMPLETED,
        )
        DocumentAsset.objects.create(
            document=document,
            asset_type=DocumentAsset.AssetType.IMAGE,
            original_name='diagram.png',
            position=0,
            ocr_status=DocumentAsset.OCRStatus.SUCCESS,
            ocr_text='识别文字',
            vision_status=DocumentAsset.VisionStatus.SUCCESS,
            vision_text='图片描述',
        )
        DocumentAsset.objects.create(
            document=document,
            asset_type=DocumentAsset.AssetType.IMAGE,
            original_name='failed.png',
            position=1,
            ocr_status=DocumentAsset.OCRStatus.FAILED,
            ocr_error='tesseract missing',
            vision_status=DocumentAsset.VisionStatus.FAILED,
            vision_error='401 unauthorized',
        )
        DocumentAsset.objects.create(
            document=document,
            asset_type=DocumentAsset.AssetType.IMAGE,
            original_name='skipped.png',
            position=2,
            ocr_status=DocumentAsset.OCRStatus.SKIPPED,
            ocr_error='No OCR text extracted',
            vision_status=DocumentAsset.VisionStatus.SKIPPED,
            vision_error='VISION_API_KEY not configured',
        )
        DocumentAsset.objects.create(
            document=document,
            asset_type=DocumentAsset.AssetType.IMAGE,
            original_name='pending.png',
            position=3,
        )

        data = DocumentSerializer(document).data

        self.assertEqual(data['asset_count'], 4)
        self.assertEqual(data['ocr_count'], 1)
        self.assertEqual(data['ocr_failed_count'], 1)
        self.assertEqual(data['ocr_skipped_count'], 1)
        self.assertEqual(data['ocr_pending_count'], 1)
        self.assertEqual(data['ocr_error_message'], 'tesseract missing')
        self.assertEqual(data['vision_count'], 1)
        self.assertEqual(data['vision_failed_count'], 1)
        self.assertEqual(data['vision_skipped_count'], 1)
        self.assertEqual(data['vision_pending_count'], 1)
        self.assertEqual(data['vision_error_message'], '401 unauthorized')
